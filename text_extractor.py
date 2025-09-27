import os
import re
from PyPDF2 import PdfReader
from docx import Document

class TextExtractor:
    """Extract text from various file formats (TXT, PDF, DOCX)"""
    
    def extract_text(self, file_path):
        """
        Extract text from a file based on its extension
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text content
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT file"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any supported encoding")
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise Exception("No text found in PDF or PDF is image-based")
            
            return self._clean_text(text)
        
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text found in DOCX file")
            
            return self._clean_text(text)
        
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with analysis
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\"\'\-\(\)]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def get_file_info(self, file_path):
        """Get basic information about the file"""
        try:
            file_size = os.path.getsize(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            return {
                'size': file_size,
                'extension': file_extension,
                'size_mb': round(file_size / (1024 * 1024), 2)
            }
        except Exception as e:
            return {'error': str(e)}